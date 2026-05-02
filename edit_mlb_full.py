"""Execute user-approved edits:
1. MLB.docx: remove F5 bet types everywhere, add concrete Pick examples in Output Format section
2. MLB workflow oTT6Iq1CPrugSDSn: re-embed amended prompt, append format compliance instruction
3. Scan other workflows for any lingering F5 references and report"""
import json, os, re, shutil, urllib.request, zipfile
from pathlib import Path
from datetime import datetime
import docx

os.environ['N8N_API_KEY'] = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJzdWIiOiJkZTM4ZjI2ZS0xZjM4LTQ1OTQtOTFhYy02ZDdlNjJkZGU3MGQiLCJpc3MiOiJuOG4iLCJhdWQiOiJwdWJsaWMtYXBpIiwiaWF0IjoxNzY5MjA1NzQ2fQ.r2CCZ2xWDdp_5RsMfXEg6cVodSrowhRh3xyOtFS0REY'

DOCX_PATH = Path(r'C:\Users\istva\OneDrive\- Dokumente -\-p\Z$$$Z\PROMPTS\MLB.docx')
BACKUP_PATH = DOCX_PATH.with_name(f'MLB_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

# ============================================================================
# STEP 1: Back up original MLB.docx
# ============================================================================
shutil.copy2(DOCX_PATH, BACKUP_PATH)
print(f"[1] Backed up original to: {BACKUP_PATH.name}")

# ============================================================================
# STEP 2: Edit MLB.docx via python-docx
# ============================================================================
doc = docx.Document(DOCX_PATH)

# We need to modify specific paragraphs:
# A) The Output Format template line: remove "/ F5 Team ML / F5 Over-Under Total"
# B) The Bet Types bullet about F5: delete entire paragraph
# C) After the Pick template line, INSERT new paragraphs with concrete examples

OUTPUT_FMT_OLD = "Pick: [Team ML / Over-Under Total / Team RL -1.5 or +1.5 / F5 Team ML / F5 Over-Under Total]"
OUTPUT_FMT_NEW = "Pick: [Team ML / Over-Under Total / Team RL -1.5 or +1.5]"

CONCRETE_EXAMPLES = [
    "Concrete examples of correctly-formatted Pick lines — substitute the template with your actual bet:",
    "- Pick: Los Angeles Dodgers ML",
    "- Pick: Under 7.5",
    "- Pick: Over 9.0",
    "- Pick: Philadelphia Phillies RL -1.5",
    "- Pick: Tampa Bay Rays RL +1.5",
]

F5_BULLET_START = "First 5 Innings (F5 ML or F5 Total)"

# Find and modify paragraphs
output_fmt_para_idx = None
f5_bullet_para_idx = None

for i, p in enumerate(doc.paragraphs):
    txt = p.text
    # A) Fix Output Format template line - remove F5 options
    if OUTPUT_FMT_OLD in txt:
        # Replace the text across runs - complicated because runs may split the string
        # Safest approach: clear all runs and write one single run with the new text
        full_new_text = txt.replace(OUTPUT_FMT_OLD, OUTPUT_FMT_NEW)
        # Keep first run's formatting, clear rest
        if p.runs:
            p.runs[0].text = full_new_text
            for r in p.runs[1:]:
                r.text = ''
        output_fmt_para_idx = i
        print(f"[2A] Modified Output Format template line (para #{i})")

    # B) Find F5 bullet for deletion
    if F5_BULLET_START in txt:
        f5_bullet_para_idx = i
        print(f"[2B] Found F5 bullet paragraph (para #{i}): {txt[:100]}...")

# Delete F5 bullet paragraph
if f5_bullet_para_idx is not None:
    p = doc.paragraphs[f5_bullet_para_idx]
    p._element.getparent().remove(p._element)
    print(f"[2B] Deleted F5 bullet paragraph")

# C) Insert concrete examples AFTER the Output Format template paragraph
# python-docx doesn't have a direct "insert paragraph after" API, so we use XML manipulation
if output_fmt_para_idx is not None:
    # Re-find the paragraph (index may have shifted after deletion)
    target_para = None
    for p in doc.paragraphs:
        if OUTPUT_FMT_NEW in p.text:
            target_para = p
            break

    if target_para is not None:
        # Insert new paragraphs after target_para
        from copy import deepcopy
        from docx.oxml.ns import qn
        parent = target_para._element.getparent()
        idx = list(parent).index(target_para._element)
        # Insert examples in reverse order (each new one inserted right after target)
        for line in reversed(CONCRETE_EXAMPLES):
            new_p = deepcopy(target_para._element)
            # Clear runs, set text
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # Create a new run with the text
            from docx.oxml import OxmlElement
            r_el = OxmlElement('w:r')
            t_el = OxmlElement('w:t')
            t_el.text = line
            t_el.set(qn('xml:space'), 'preserve')
            r_el.append(t_el)
            new_p.append(r_el)
            parent.insert(idx + 1, new_p)
        print(f"[2C] Inserted {len(CONCRETE_EXAMPLES)} concrete example paragraphs after Output Format template")

doc.save(DOCX_PATH)
print(f"[2] Saved modified MLB.docx")

# ============================================================================
# STEP 3: Re-extract and verify the amended prompt
# ============================================================================
def extract_docx_text(path):
    with zipfile.ZipFile(path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    text = xml
    text = re.sub(r'<w:br\s*/>', '\n', text)
    text = re.sub(r'<w:tab\s*/>', '\t', text)
    text = re.sub(r'</w:p>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"').replace('&apos;', "'")
    lines = [l.rstrip() for l in text.split('\n')]
    while lines and not lines[0].strip(): lines.pop(0)
    while lines and not lines[-1].strip(): lines.pop()
    return '\n'.join(lines)

amended_text = extract_docx_text(DOCX_PATH)
print(f"\n[3] Re-extracted amended MLB.docx: {len(amended_text)} chars")

# Fidelity checks
assert "F5 Team ML" not in amended_text, "F5 Team ML still present!"
assert "F5 Over-Under" not in amended_text, "F5 Over-Under still present!"
assert "First 5 Innings" not in amended_text, "F5 bullet still present!"
assert "Pick: Under 7.5" in amended_text, "Concrete examples not added!"
assert "Pick: Over 9.0" in amended_text, "Concrete examples not added!"
print("[3] Fidelity check passed: F5 removed, concrete examples added")

# Save extracted text
out_dir = Path(r'C:\Users\istva\.claude\CODE\EDGE STACKER')
(out_dir / 'MLB_AMENDED_v3.txt').write_text(amended_text, encoding='utf-8')

# ============================================================================
# STEP 4: Update MLB workflow (oTT6Iq1CPrugSDSn)
# ============================================================================
WORKFLOW_ID = 'oTT6Iq1CPrugSDSn'

# Fetch workflow
req = urllib.request.Request(
    f"http://localhost:5678/api/v1/workflows/{WORKFLOW_ID}",
    headers={"X-N8N-API-KEY": os.environ['N8N_API_KEY']}
)
with urllib.request.urlopen(req, timeout=30) as resp:
    wf = json.loads(resp.read().decode('utf-8'))

# Patch Build Gemini Request node
build_node = next(n for n in wf['nodes'] if n['name'] == 'Build Gemini Request')
current_code = build_node['parameters']['jsCode']

# Replace embedded prompt
# Find existing systemPrompt assignment and replace it
def new_code_with_amended_prompt(old_code, new_prompt_text):
    # Pattern: const systemPrompt = "..."
    pattern = r'(const systemPrompt = )(".*?");'
    replacement = lambda m: m.group(1) + json.dumps(new_prompt_text) + ';'
    return re.sub(pattern, replacement, old_code, count=1, flags=re.DOTALL)

new_code = new_code_with_amended_prompt(current_code, amended_text)

# Append format-compliance instruction to userMessage
# Current user message ends with: "First line MUST be: MLB Slate — ${dateShort} — Top Plays`;
FORMAT_INSTRUCTION = "\\n\\nIMPORTANT FORMAT COMPLIANCE: The Pick line MUST contain the CONCRETE chosen bet, never a template placeholder. For Totals write \\\"Pick: Over 7.5\\\" or \\\"Pick: Under 8.0\\\" with the actual number. For moneylines write \\\"Pick: Los Angeles Dodgers ML\\\". For run lines write \\\"Pick: Philadelphia Phillies RL -1.5\\\". The square brackets in the system prompt's Output Format are option indicators; you must substitute them with your actual chosen bet including the specific number/line."

old_suffix = "First line MUST be: MLB Slate \u2014 ${dateShort} \u2014 Top Plays`;"
new_suffix = f"First line MUST be: MLB Slate \u2014 ${{dateShort}} \u2014 Top Plays{FORMAT_INSTRUCTION}`;"

assert old_suffix in new_code, f"Could not find user message suffix to amend"
new_code = new_code.replace(old_suffix, new_suffix)

build_node['parameters']['jsCode'] = new_code

# PUT update
body = {"name": wf["name"], "nodes": wf["nodes"], "connections": wf["connections"], "settings": wf["settings"]}
data = json.dumps(body, ensure_ascii=False).encode('utf-8')
req = urllib.request.Request(
    f"http://localhost:5678/api/v1/workflows/{WORKFLOW_ID}",
    data=data,
    headers={"Content-Type": "application/json", "X-N8N-API-KEY": os.environ['N8N_API_KEY']},
    method="PUT"
)
with urllib.request.urlopen(req, timeout=30) as resp:
    r = json.loads(resp.read().decode('utf-8'))
    print(f"\n[4] Updated MLB workflow '{r['name']}':")
    print(f"    - Embedded prompt replaced with amended docx text ({len(amended_text)} chars, F5 removed)")
    print(f"    - Format-compliance instruction appended to user message")

# ============================================================================
# STEP 5: Scan ALL other workflows for F5 references
# ============================================================================
print(f"\n[5] Scanning all other workflows for F5 references...")
req = urllib.request.Request(
    "http://localhost:5678/api/v1/workflows",
    headers={"X-N8N-API-KEY": os.environ['N8N_API_KEY']}
)
with urllib.request.urlopen(req, timeout=30) as resp:
    all_wfs = json.loads(resp.read().decode('utf-8'))

f5_patterns = [r'\bF5 ML\b', r'\bF5 Team ML\b', r'\bF5 Over[- ]Under\b', r'\bF5 Total\b', r'First 5 Innings', r'\(F5 ML']
other_wfs_with_f5 = []
for summary in all_wfs.get('data', []):
    wf_id = summary['id']
    if wf_id == WORKFLOW_ID:
        continue
    req = urllib.request.Request(
        f"http://localhost:5678/api/v1/workflows/{wf_id}",
        headers={"X-N8N-API-KEY": os.environ['N8N_API_KEY']}
    )
    with urllib.request.urlopen(req, timeout=30) as r:
        full = json.loads(r.read().decode('utf-8'))
    full_str = json.dumps(full)
    hits = [p for p in f5_patterns if re.search(p, full_str)]
    if hits:
        other_wfs_with_f5.append((wf_id, summary['name'], hits))

if other_wfs_with_f5:
    print(f"    F5 references found in {len(other_wfs_with_f5)} other workflow(s):")
    for wf_id, name, hits in other_wfs_with_f5:
        print(f"      - {name} ({wf_id}): matches {hits}")
else:
    print("    No F5 references found in any other workflow.")
